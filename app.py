"""
Sistema de Visitas Técnicas - CVP (Caja de Vivienda Popular)
Bogotá, Colombia
"""

import streamlit as st
import pandas as pd
import os
import io
from pathlib import Path
from datetime import date, datetime, time
import plotly.express as px
import plotly.graph_objects as go

# ─────────────────────────────────────────────
# GLOBAL CONFIG
# ─────────────────────────────────────────────
DATA_DIR     = Path(__file__).parent / "data"
FORMATOS_DIR = Path(__file__).parent / "formatos"
MAESTRO_PATH    = DATA_DIR / "maestro_predios_cvp.csv"
VISITAS_PATH    = DATA_DIR / "visitas.csv"
RESULTADOS_PATH = DATA_DIR / "resultados.csv"
TECNICOS_PATH   = DATA_DIR / "tecnicos.csv"
FICHA_TEMPLATE  = FORMATOS_DIR / "208-REAS-Ft-30 FICHA TECNICA DE RECONOCIMIENTO v7.xlsx"
INFORME_TEMPLATE= FORMATOS_DIR / "208-REAS-Ft-176 INFORME DE GESTION.docx"

COLOR_PRIMARY   = "#CC0000"   # Rojo CVP
COLOR_SECONDARY = "#007A3D"   # Verde CVP
COLOR_AMARILLO  = "#FFC300"   # Amarillo CVP
COLOR_SUCCESS   = "#007A3D"
COLOR_DANGER    = "#CC0000"
COLOR_WARNING   = "#FFC300"

CARGOS = [
    "Técnico Social",
    "Profesional Social",
    "Técnico Jurídico",
    "Profesional Jurídico",
    "Técnico Predial",
]

ESTADOS_VISITA = ["Pendiente", "Exitosa", "Fallida"]
OCUPACIONES = ["Ocupado", "Desocupado", "En construcción", "Lote vacío"]
TIPOS_CONSTRUCCION = [
    "Mampostería confinada",
    "Mampostería simple",
    "Madera",
    "Metálica",
    "Mixta",
    "Otro",
]
ESTADOS_CONSERVACION = ["Bueno", "Regular", "Malo", "Ruina"]
PISOS_OPTIONS = ["1", "2", "3", "4", "5", "6+"]
MOTIVOS_FALLIDA = [
    "Predio no encontrado",
    "Acceso negado",
    "Propietario no presente",
    "Predio demolido",
    "Dirección incorrecta",
    "Otro",
]

st.set_page_config(
    page_title="CVP - Sistema de Visitas",
    page_icon="🏠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
# CUSTOM CSS
# ─────────────────────────────────────────────
st.markdown(
    f"""
    <style>
    /* Sidebar background */
    [data-testid="stSidebar"] {{
        background-color: {COLOR_PRIMARY};
        border-right: 4px solid {COLOR_AMARILLO};
    }}
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div,
    [data-testid="stSidebar"] label {{
        color: white !important;
    }}
    /* Header strip */
    .cvp-header {{
        background-color: {COLOR_PRIMARY};
        color: white;
        padding: 18px 24px;
        border-radius: 8px;
        margin-bottom: 20px;
    }}
    .cvp-header h1 {{
        margin: 0;
        font-size: 1.6rem;
    }}
    .cvp-header p {{
        margin: 4px 0 0;
        font-size: 0.9rem;
        opacity: 0.85;
    }}
    /* Metric card */
    .metric-card {{
        background: white;
        border-left: 5px solid {COLOR_SECONDARY};
        border-radius: 6px;
        padding: 16px 20px;
        box-shadow: 0 2px 6px rgba(0,0,0,0.08);
        text-align: center;
    }}
    .metric-card .value {{
        font-size: 2.2rem;
        font-weight: 700;
        color: {COLOR_PRIMARY};
    }}
    .metric-card .label {{
        font-size: 0.85rem;
        color: #555;
        margin-top: 4px;
    }}
    /* Section title */
    .section-title {{
        color: {COLOR_PRIMARY};
        font-weight: 700;
        font-size: 1.1rem;
        border-bottom: 3px solid {COLOR_SECONDARY};
        padding-bottom: 4px;
        margin: 20px 0 12px;
    }}
    /* Result banner */
    .banner-exitosa {{
        background-color: {COLOR_SUCCESS};
        color: white;
        padding: 10px 20px;
        border-radius: 6px;
        font-weight: 700;
        font-size: 1.1rem;
        text-align: center;
    }}
    .banner-fallida {{
        background-color: {COLOR_DANGER};
        color: white;
        padding: 10px 20px;
        border-radius: 6px;
        font-weight: 700;
        font-size: 1.1rem;
        text-align: center;
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────
# DATA LOADING HELPERS
# ─────────────────────────────────────────────

@st.cache_resource(show_spinner=False)
def get_supabase():
    try:
        from supabase import create_client
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except Exception:
        return None


@st.cache_data(show_spinner=False)
def load_maestro():
    try:
        df = pd.read_csv(MAESTRO_PATH, encoding="utf-8-sig", low_memory=False, dtype=str)
        df.columns = df.columns.str.strip()
        for col in ["LATITUD", "LONGITUD"]:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")
        if "AVALUO" in df.columns:
            df["AVALUO"] = pd.to_numeric(df["AVALUO"], errors="coerce").fillna(0)
        return df
    except Exception as e:
        st.error(f"Error cargando maestro: {e}")
        return pd.DataFrame()


def load_visitas():
    sb = get_supabase()
    if sb:
        try:
            resp = sb.table("visitas").select("*").execute()
            if resp.data:
                df = pd.DataFrame(resp.data)
                df = df.drop(columns=["id"], errors="ignore")
                df = df.astype(str).replace("None", "").replace("nan", "")
                return df
        except Exception as e:
            st.warning(f"Supabase (visitas): {e}")
    # Fallback CSV
    try:
        df = pd.read_csv(VISITAS_PATH, dtype=str)
        if df.empty or len(df.columns) == 0:
            return _empty_visitas()
        df.columns = df.columns.str.strip()
        return df
    except Exception:
        return _empty_visitas()


def _empty_visitas():
    return pd.DataFrame(
        columns=[
            "NUM_VISITA", "FECHA_PROGRAMADA", "REA", "SIN_REA",
            "DIRECCION_MANUAL", "LATITUD_MANUAL", "LONGITUD_MANUAL",
            "TECNICOS", "ESTADO", "NUM_VISITA_PREDIO", "FECHA_REGISTRO",
            "OBSERVACIONES_PROG",
        ]
    )


def load_resultados():
    sb = get_supabase()
    if sb:
        try:
            resp = sb.table("resultados").select("*").execute()
            if resp.data:
                df = pd.DataFrame(resp.data)
                df = df.drop(columns=["id"], errors="ignore")
                df = df.astype(str).replace("None", "").replace("nan", "")
                return df
        except Exception as e:
            st.warning(f"Supabase (resultados): {e}")
    # Fallback CSV
    try:
        df = pd.read_csv(RESULTADOS_PATH, dtype=str)
        if df.empty or len(df.columns) == 0:
            return _empty_resultados()
        df.columns = df.columns.str.strip()
        return df
    except Exception:
        return _empty_resultados()


def _empty_resultados():
    return pd.DataFrame(
        columns=[
            "NUM_VISITA", "REA", "FECHA_VISITA", "HORA_INICIO", "HORA_FIN",
            "TECNICOS", "RESULTADO", "OCUPACION", "PROP_CONTACTADO",
            "TIPO_CONSTRUCCION", "NUM_PISOS", "ESTADO_CONSERVACION",
            "LINDERO_NORTE", "LINDERO_SUR", "LINDERO_ORIENTE", "LINDERO_OCCIDENTE",
            "TIPO_INMUEBLE", "ESTRATO", "UNIDADES_VIVIENDA",
            "UPL", "UPZ", "AREA_TERRENO", "AREA_CONSTRUCCION",
            "TIPO_GESTION", "TELEFONO_BENEFICIARIO", "CORREO_BENEFICIARIO", "COMPONENTE",
            "MOTIVO_FALLIDA", "OBSERVACIONES", "FOTOS", "FECHA_REGISTRO",
        ]
    )


def load_tecnicos():
    try:
        df = pd.read_csv(TECNICOS_PATH, dtype=str)
        df.columns = df.columns.str.strip()
        df["ACTIVO"] = df["ACTIVO"].str.strip().str.lower().map(
            {"true": True, "false": False, "1": True, "0": False}
        ).fillna(True)
        return df
    except Exception:
        return pd.DataFrame(
            columns=["ID_TECNICO", "NOMBRE", "CARGO", "EMAIL", "ACTIVO"]
        )


def save_visitas(df):
    df.to_csv(VISITAS_PATH, index=False)
    sb = get_supabase()
    if sb and not df.empty:
        try:
            records = df.where(pd.notna(df), None).to_dict("records")
            sb.table("visitas").upsert(records, on_conflict="NUM_VISITA").execute()
        except Exception as e:
            st.warning(f"No se pudo sincronizar visitas con Supabase: {e}")


def save_resultados(df):
    df.to_csv(RESULTADOS_PATH, index=False)
    sb = get_supabase()
    if sb and not df.empty:
        try:
            records = df.where(pd.notna(df), None).to_dict("records")
            sb.table("resultados").upsert(records, on_conflict="NUM_VISITA").execute()
        except Exception as e:
            st.warning(f"No se pudo sincronizar resultados con Supabase: {e}")


def delete_visita_supabase(num_visita: str):
    sb = get_supabase()
    if sb:
        try:
            sb.table("visitas").delete().eq("NUM_VISITA", num_visita).execute()
        except Exception as e:
            st.warning(f"No se pudo eliminar de Supabase: {e}")


def save_tecnicos(df):
    df.to_csv(TECNICOS_PATH, index=False)


def next_num_visita(visitas_df):
    year = datetime.now().year
    prefix = f"VT-{year}-"
    existing = visitas_df["NUM_VISITA"].dropna() if "NUM_VISITA" in visitas_df.columns else pd.Series(dtype=str)
    year_nums = existing[existing.str.startswith(prefix, na=False)]
    if year_nums.empty:
        return f"{prefix}0001"
    maxn = year_nums.str.replace(prefix, "", regex=False).apply(
        lambda x: int(x) if x.isdigit() else 0
    ).max()
    return f"{prefix}{str(maxn + 1).zfill(4)}"


def count_visitas_predio(visitas_df, rea):
    if visitas_df.empty or "REA" not in visitas_df.columns:
        return 0
    return int((visitas_df["REA"] == str(rea)).sum())


# ─────────────────────────────────────────────
# SIDEBAR NAVIGATION
# ─────────────────────────────────────────────

with st.sidebar:
    st.markdown(
        f"""
        <div style='text-align:center; padding: 10px 0 20px;'>
            <div style='font-size:1.3rem; font-weight:700; color:white; letter-spacing:2px;'>CVP</div>
            <div style='font-size:0.75rem; color:#aac4e0;'>Caja de Vivienda Popular<br>Bogotá</div>
        </div>
        <hr style='border-color: #336699; margin: 0 0 12px;'>
        """,
        unsafe_allow_html=True,
    )

    pagina = st.radio(
        "Navegación",
        options=[
            "Inicio",
            "Programar Visita",
            "Visitas Programadas",
            "Registrar Resultado",
            "Gestión de Técnicos",
            "Indicadores",
        ],
        label_visibility="collapsed",
    )

    st.markdown(
        """
        <hr style='border-color: #336699; margin: 20px 0 10px;'>
        <div style='font-size:0.7rem; color:#aac4e0; text-align:center;'>
            Dirección de Reasentamientos<br>© 2025 CVP Bogotá
        </div>
        """,
        unsafe_allow_html=True,
    )

# ─────────────────────────────────────────────
# SESSION STATE INIT
# ─────────────────────────────────────────────
if "pdf_num_visita" not in st.session_state:
    st.session_state["pdf_num_visita"] = None
if "registro_num_visita" not in st.session_state:
    st.session_state["registro_num_visita"] = None
if "mostrar_descarga" not in st.session_state:
    st.session_state["mostrar_descarga"] = None


# ════════════════════════════════════════════════════════════════
# FORMAT GENERATION FUNCTIONS
# ════════════════════════════════════════════════════════════════

def _str(val):
    """Return clean string or empty."""
    if val is None or (isinstance(val, float) and str(val) == "nan"):
        return ""
    return str(val).strip()


def generar_ficha_tecnica(res_row, vis_row, maestro_row):
    """Pre-fill 208-REAS-Ft-30 FICHA TECNICA DE RECONOCIMIENTO v7.xlsx and return bytes."""
    import openpyxl
    from copy import copy

    wb = openpyxl.load_workbook(str(FICHA_TEMPLATE))
    ws = wb.active

    fecha_vis  = _str(res_row.get("FECHA_VISITA", ""))
    rea        = _str(res_row.get("REA", ""))
    tecnicos   = _str(res_row.get("TECNICOS", "")).replace("|", " / ")
    localidad  = _str(maestro_row.get("LOCALIDAD", ""))  if not maestro_row.empty else ""
    barrio     = _str(maestro_row.get("BARRIO", ""))     if not maestro_row.empty else ""
    dir_campo  = _str(maestro_row.get("DIRECCION", ""))  if not maestro_row.empty else ""
    manzana    = _str(maestro_row.get("MANZANA", ""))    if not maestro_row.empty else ""
    lote       = _str(maestro_row.get("LOTE", ""))       if not maestro_row.empty else ""
    chip       = _str(maestro_row.get("CHIP", ""))       if not maestro_row.empty else ""
    latitud    = _str(maestro_row.get("LATITUD", ""))    if not maestro_row.empty else ""
    longitud   = _str(maestro_row.get("LONGITUD", ""))   if not maestro_row.empty else ""
    lind_n     = _str(res_row.get("LINDERO_NORTE", ""))
    lind_s     = _str(res_row.get("LINDERO_SUR", ""))
    lind_or    = _str(res_row.get("LINDERO_ORIENTE", ""))
    lind_oc    = _str(res_row.get("LINDERO_OCCIDENTE", ""))

    # Cargo del técnico desde tecnicos.csv
    cargo_tec = ""
    try:
        tec_path = Path(__file__).parent / "data" / "tecnicos.csv"
        tec_df_f = pd.read_csv(tec_path, encoding="utf-8-sig", dtype=str)
        primer_tec = tecnicos.split("/")[0].strip()
        tec_info_f = tec_df_f[tec_df_f["NOMBRE"].str.strip() == primer_tec]
        if not tec_info_f.empty and "CARGO" in tec_info_f.columns:
            cargo_tec = _str(tec_info_f["CARGO"].iloc[0])
    except Exception:
        cargo_tec = ""

    # Cell mapping  (label row → value row, first cell of merged range)
    ws["A7"]  = fecha_vis       # FECHA DE ELABORACION
    ws["N7"]  = rea             # IDENTIFICADOR
    # A9 = NOMBRE DE QUIEN ATIENDE LA VISITA → se deja en blanco (diligenciar a mano con nombre y CC)
    ws["A12"] = localidad       # 2.1 LOCALIDAD
    ws["K12"] = barrio          # 2.2 BARRIO
    ws["A14"] = dir_campo       # 2.4 DIRECCION TOMADA EN CAMPO
    upl_val = _str(res_row.get("UPL", "")) or (_str(maestro_row.get("UPL", "")) if not maestro_row.empty else "")
    upz_val = _str(res_row.get("UPZ", "")) or (_str(maestro_row.get("UPZ", "")) if not maestro_row.empty else "")
    ws["S14"] = upl_val                                   # 2.5 UPL
    ws["W14"] = upz_val                                   # 2.6 UPZ
    ws["A18"] = dir_campo       # 2.8 DIRECCIÓN CATASTRAL
    ws["S18"] = manzana         # MANZANA catastral
    ws["W18"] = lote            # LOTE catastral
    ws["A20"] = chip            # 2.9 CHIP CATASTRAL
    ws["K20"] = _str(res_row.get("AREA_TERRENO", ""))    # 2.10 ÁREA TERRENO
    ws["S20"] = _str(res_row.get("AREA_CONSTRUCCION", "")) # 2.11 ÁREA CONSTRUCCIÓN
    ws["Z20"] = _str(res_row.get("ESTRATO", ""))         # 2.12 ESTRATO
    ws["AE12"] = latitud        # 2.14 COORDENADA X (valor junto a "X:")
    ws["AR12"] = longitud       # 2.14 COORDENADA Y (valor junto a "Y:")
    ws["K28"] = _str(res_row.get("TIPO_INMUEBLE", ""))   # 3.1 TIPO INMUEBLE
    ws["N29"] = _str(res_row.get("NUM_PISOS", ""))        # 3.2 NÚMERO DE PISOS
    ws["S29"] = _str(res_row.get("UNIDADES_VIVIENDA", "")) # 3.3 UNIDADES VIVIENDA
    ws["F22"] = lind_n          # LINDERO NORTE
    ws["F23"] = lind_s          # LINDERO SUR
    ws["F24"] = lind_or         # LINDERO ORIENTE
    ws["F25"] = lind_oc         # LINDERO OCCIDENTE
    ws["B97"] = f"NOMBRE        {tecnicos}"   # ELABORÓ - nombre técnico
    ws["B98"] = f"CARGO          {cargo_tec}" # ELABORÓ - cargo técnico

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def generar_informe_gestion(res_row, vis_row, maestro_row):
    """Pre-fill 208-REAS-Ft-176 INFORME DE GESTION.docx and return bytes."""
    from docx import Document

    doc = Document(str(INFORME_TEMPLATE))

    rea         = _str(res_row.get("REA", ""))
    propietario = _str(maestro_row.get("PROPIETARIO_1", "")) if not maestro_row.empty else ""
    tecnicos    = _str(res_row.get("TECNICOS", "")).replace("|", " / ")
    fecha_vis   = _str(res_row.get("FECHA_VISITA", ""))
    motivo      = _str(res_row.get("MOTIVO_FALLIDA", "")).replace("|", ", ")
    observ      = _str(res_row.get("OBSERVACIONES", ""))
    telefono    = _str(res_row.get("TELEFONO_BENEFICIARIO", ""))
    correo      = _str(res_row.get("CORREO_BENEFICIARIO", ""))
    componente  = _str(res_row.get("COMPONENTE", "")) or "Reasentamiento"
    tipo_gestion = _str(res_row.get("TIPO_GESTION", "")) or "Visita a campo"

    # Table 0: datos básicos
    tabla = doc.tables[0]
    tabla.rows[0].cells[1].text = rea           # IDENTIFICADOR
    tabla.rows[1].cells[1].text = propietario   # BENEFICIARIO
    tabla.rows[2].cells[1].text = telefono      # TELEFONO
    tabla.rows[3].cells[1].text = correo        # CORREO ELECTRÓNICO
    tabla.rows[4].cells[1].text = tecnicos      # NOMBRE PROFESIONAL
    tabla.rows[5].cells[1].text = componente    # COMPONENTE
    tabla.rows[6].cells[1].text = fecha_vis     # FECHA

    # Table 1: firma — pre-llenar nombre y contrato del técnico
    # Cargar datos del técnico desde tecnicos.csv para obtener contrato
    try:
        from pathlib import Path
        import pandas as pd
        tec_path = Path(__file__).parent / "data" / "tecnicos.csv"
        tec_df = pd.read_csv(tec_path, encoding="utf-8-sig", dtype=str)
        primer_tec = tecnicos.split("/")[0].strip()
        tec_info = tec_df[tec_df["NOMBRE"].str.strip() == primer_tec]
        contrato = tec_info["CONTRATO"].iloc[0].strip() if not tec_info.empty and "CONTRATO" in tec_info.columns and pd.notna(tec_info["CONTRATO"].iloc[0]) else ""
    except Exception:
        contrato = ""

    firma_table = doc.tables[1]
    # Cell [0,0]: firma + nombre técnico
    firma_table.rows[0].cells[0].paragraphs[-1].clear()
    firma_table.rows[0].cells[0].paragraphs[-1].add_run(tecnicos)
    # Cell [0,1]: contrato
    if contrato:
        firma_table.rows[0].cells[1].paragraphs[-1].clear()
        firma_table.rows[0].cells[1].paragraphs[-1].add_run(contrato)

    # Marcar checkbox en P01 según tipo_gestion
    # Formato original: (       )  Seguimiento al proceso    (        ) Visita a campo       (       ) Entrega de documentos
    paragraphs = doc.paragraphs
    for p in paragraphs:
        if "Seguimiento al proceso" in p.text and "Visita a campo" in p.text:
            new_text = p.text
            opciones = ["Seguimiento al proceso", "Visita a campo", "Entrega de documentos"]
            for opcion in opciones:
                if opcion == tipo_gestion:
                    new_text = new_text.replace("(       )", "(  X  )", 1)
                else:
                    idx_op = new_text.find(opcion)
                    if idx_op > 0:
                        paren_start = new_text.rfind("(", 0, idx_op)
                        paren_end = new_text.find(")", paren_start)
                        if paren_start >= 0 and paren_end > paren_start:
                            new_text = new_text[:paren_start] + "(       )" + new_text[paren_end + 1:]
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            break

    # Escribir en el párrafo vacío SIGUIENTE a "Descripción Gestión realizada:" (P05)
    desc_text = ""
    if motivo:
        desc_text += "Motivo de visita fallida: " + motivo
    if observ:
        if desc_text:
            desc_text += "\n\n"
        desc_text += "Descripción de la gestión realizada: " + observ
    for idx, p in enumerate(paragraphs):
        if "Descripci" in p.text and "Gesti" in p.text:
            siguiente = idx + 1
            if siguiente < len(paragraphs):
                target = paragraphs[siguiente]
                target.clear()
                target.add_run(desc_text)
            break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ════════════════════════════════════════════════════════════════
# PAGE: INICIO
# ════════════════════════════════════════════════════════════════
if pagina == "Inicio":
    st.markdown(
        f"""
        <div class="cvp-header">
            <h1>🏛️ Sistema de Visitas Técnicas — CVP</h1>
            <p>Caja de Vivienda Popular · Dirección de Reasentamientos · Bogotá D.C.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    maestro    = load_maestro()
    visitas    = load_visitas()
    resultados = load_resultados()

    total_predios     = len(maestro)
    total_programadas = len(visitas) if not visitas.empty else 0

    if not visitas.empty and "ESTADO" in visitas.columns:
        total_exitosas   = int((visitas["ESTADO"] == "Exitosa").sum())
        total_fallidas   = int((visitas["ESTADO"] == "Fallida").sum())
        total_pendientes = int((visitas["ESTADO"] == "Pendiente").sum())
    else:
        total_exitosas = total_fallidas = total_pendientes = 0

    total_realizadas = total_exitosas + total_fallidas
    pct_exito = round(total_exitosas / total_realizadas * 100, 1) if total_realizadas > 0 else 0

    rend_semanal = 0.0
    if not visitas.empty and "FECHA_PROGRAMADA" in visitas.columns:
        v_rend = visitas.copy()
        v_rend["_fd"] = pd.to_datetime(v_rend["FECHA_PROGRAMADA"], errors="coerce")
        cutoff = pd.Timestamp(date.today()) - pd.Timedelta(weeks=4)
        v_rend = v_rend[v_rend["_fd"] >= cutoff]
        if not v_rend.empty:
            semanas = v_rend["_fd"].dt.to_period("W").nunique()
            rend_semanal = round(len(v_rend) / max(semanas, 1), 1)

    c1, c2, c3, c4 = st.columns(4)
    kpis_r1 = [
        (c1, total_predios,     "Total Predios",        "🗂️",  COLOR_PRIMARY),
        (c2, total_programadas, "Visitas Programadas",  "📋",  "#1565C0"),
        (c3, total_realizadas,  "Visitas Realizadas",   "✅",  COLOR_SUCCESS),
        (c4, total_pendientes,  "Visitas Pendientes",   "⏳",  "#E65100"),
    ]
    for col, val, label, icon, color in kpis_r1:
        with col:
            st.markdown(
                f"""<div style="background:white;border-left:5px solid {color};border-radius:8px;
                    padding:18px 20px;box-shadow:0 2px 8px rgba(0,0,0,0.09);text-align:center;margin-bottom:4px;">
                    <div style="font-size:1.6rem;">{icon}</div>
                    <div style="font-size:2rem;font-weight:800;color:{color};">{val:,}</div>
                    <div style="font-size:0.82rem;color:#555;margin-top:2px;">{label}</div>
                </div>""",
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)

    c5, c6, c7, c8 = st.columns(4)
    kpis_r2 = [
        (c5, total_exitosas,  "Visitas Exitosas",              "🟢", COLOR_SUCCESS),
        (c6, total_fallidas,  "Visitas Fallidas",              "🔴", COLOR_DANGER),
        (c7, f"{pct_exito}%", "Porcentaje de Éxito",          "🎯", "#6A1B9A"),
        (c8, rend_semanal,    "Prom. Visitas/Semana (1 mes)", "📈", "#00695C"),
    ]
    for col, val, label, icon, color in kpis_r2:
        with col:
            st.markdown(
                f"""<div style="background:white;border-left:5px solid {color};border-radius:8px;
                    padding:18px 20px;box-shadow:0 2px 8px rgba(0,0,0,0.09);text-align:center;margin-bottom:4px;">
                    <div style="font-size:1.6rem;">{icon}</div>
                    <div style="font-size:2rem;font-weight:800;color:{color};">{val}</div>
                    <div style="font-size:0.82rem;color:#555;margin-top:2px;">{label}</div>
                </div>""",
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)

    if not visitas.empty and "FECHA_PROGRAMADA" in visitas.columns and "ESTADO" in visitas.columns:
        v_spark = visitas.copy()
        v_spark["_fd"] = pd.to_datetime(v_spark["FECHA_PROGRAMADA"], errors="coerce")
        cutoff8 = pd.Timestamp(date.today()) - pd.Timedelta(weeks=8)
        v_spark = v_spark[v_spark["_fd"] >= cutoff8].copy()
        if not v_spark.empty:
            v_spark["_sem"] = v_spark["_fd"].dt.to_period("W").apply(
                lambda r: str(r.start_time.date()) if pd.notna(r) else None
            )
            spark_data = v_spark.groupby(["_sem", "ESTADO"]).size().reset_index(name="n")
            spark_data = spark_data.dropna(subset=["_sem"]).sort_values("_sem")
            fig_spark = go.Figure()
            for estado, color in [("Exitosa", COLOR_SUCCESS), ("Fallida", COLOR_DANGER), ("Pendiente", COLOR_WARNING)]:
                d = spark_data[spark_data["ESTADO"] == estado]
                if not d.empty:
                    fig_spark.add_trace(go.Bar(
                        x=d["_sem"], y=d["n"], name=estado,
                        marker_color=color, marker_line_width=0,
                        hovertemplate="%{x}<br>%{y} visitas<extra></extra>",
                    ))
            fig_spark.update_layout(
                barmode="stack",
                title=dict(text="Evolución de visitas — últimas 8 semanas", font=dict(size=14, color="#333")),
                height=250, margin=dict(l=20, r=20, t=40, b=20),
                plot_bgcolor="white", paper_bgcolor="white",
                legend=dict(orientation="h", yanchor="bottom", y=1.01, xanchor="right", x=1),
                xaxis=dict(showgrid=False, tickfont=dict(size=10)),
                yaxis=dict(showgrid=True, gridcolor="#f0f0f0", tickfont=dict(size=10)),
            )
            st.plotly_chart(fig_spark, use_container_width=True)

    st.markdown('<div class="section-title">Acciones rápidas</div>', unsafe_allow_html=True)
    qcol1, qcol2, qcol3 = st.columns(3)
    with qcol1:
        st.info("📋 **Programar Visita**\nRegistre una nueva visita técnica a un predio del maestro.")
    with qcol2:
        st.info("✅ **Registrar Resultado**\nDocumente el resultado y descargue el formato oficial.")
    with qcol3:
        st.info("📊 **Indicadores**\nConsulte métricas, reportes Excel y análisis por técnico.")


# ════════════════════════════════════════════════════════════════
# PAGE: PROGRAMAR VISITA
# ════════════════════════════════════════════════════════════════
elif pagina == "Programar Visita":
    st.markdown(
        '<div class="cvp-header"><h1>Programar Visita Técnica</h1><p>Registre y programe visitas a predios del reasentamiento</p></div>',
        unsafe_allow_html=True,
    )

    maestro = load_maestro()
    tecnicos_df = load_tecnicos()
    tecnicos_activos = (
        tecnicos_df[tecnicos_df["ACTIVO"] == True]["NOMBRE"].tolist()
        if not tecnicos_df.empty else []
    )

    tab_rea, tab_sin_rea = st.tabs(["🔎 Con REA", "📍 Sin REA"])

    # ── TAB CON REA ──────────────────────────────────────────────
    with tab_rea:
        modo = st.radio("Modo de ingreso", ["Manual", "Carga masiva Excel"], horizontal=True)

        predios_seleccionados = []

        if modo == "Manual":
            # ── Inicializar lista acumulable en session_state ──────────
            if "lista_predios_manual" not in st.session_state:
                st.session_state["lista_predios_manual"] = []

            st.markdown('<div class="section-title">Agregar predios a la visita</div>', unsafe_allow_html=True)

            campo_busqueda = st.radio("Buscar por", ["REA", "CHIP"], horizontal=True, key="campo_busq_manual")

            if campo_busqueda == "REA":
                opciones = sorted(maestro["REA"].dropna().astype(str).str.strip().unique().tolist())
            else:
                opciones = sorted(maestro["CHIP"].dropna().astype(str).str.strip().unique().tolist()) if "CHIP" in maestro.columns else []

            col_sel, col_btn = st.columns([4, 1])
            with col_sel:
                valor_busqueda = st.selectbox(
                    f"Escriba o seleccione {campo_busqueda}",
                    options=[""] + opciones,
                    index=0,
                    placeholder=f"Empiece a escribir el {campo_busqueda}...",
                    help=f"Escriba los primeros dígitos para filtrar ({len(opciones):,} opciones)",
                    key="sel_busq_manual",
                )
            with col_btn:
                st.markdown("<br>", unsafe_allow_html=True)
                agregar = st.button("Agregar", use_container_width=True)

            # ── Agregar predio a la lista ──────────────────────────────
            if agregar and valor_busqueda:
                rea_a_agregar = valor_busqueda.strip()
                if campo_busqueda == "CHIP":
                    match = maestro[maestro["CHIP"].str.strip() == rea_a_agregar] if "CHIP" in maestro.columns else pd.DataFrame()
                    rea_a_agregar = match["REA"].iloc[0] if not match.empty and pd.notna(match["REA"].iloc[0]) else rea_a_agregar
                if rea_a_agregar in st.session_state["lista_predios_manual"]:
                    st.warning(f"El predio **{rea_a_agregar}** ya está en la lista.")
                else:
                    st.session_state["lista_predios_manual"].append(rea_a_agregar)
                    st.success(f"✓ Predio **{rea_a_agregar}** agregado.")

            # ── Mostrar lista acumulada ────────────────────────────────
            lista_actual = st.session_state["lista_predios_manual"]
            if lista_actual:
                st.markdown(f'<div class="section-title">Predios en esta visita ({len(lista_actual)})</div>', unsafe_allow_html=True)

                show_cols = [c for c in ["REA", "CHIP", "CHIP_VALIDADO", "LOCALIDAD", "BARRIO", "DIRECCION", "ESTADO_REA", "ENLACE_MAPS"] if c in maestro.columns]
                df_lista = maestro[maestro["REA"].isin(lista_actual)][show_cols].copy()

                # Agregar predios sin REA en maestro (sin REA) si los hay
                reas_no_en_maestro = [r for r in lista_actual if r not in maestro["REA"].values]
                if reas_no_en_maestro:
                    extras = pd.DataFrame([{"REA": r} for r in reas_no_en_maestro])
                    df_lista = pd.concat([df_lista, extras], ignore_index=True)

                st.dataframe(
                    df_lista,
                    use_container_width=True,
                    hide_index=True,
                    column_config={
                        "ENLACE_MAPS": st.column_config.LinkColumn("Maps", display_text="Ver mapa"),
                    },
                )

                # Eliminar predio individual
                col_rm1, col_rm2 = st.columns([3, 1])
                with col_rm1:
                    predio_quitar = st.selectbox("Quitar predio de la lista", options=[""] + lista_actual, key="sel_quitar")
                with col_rm2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("Quitar", use_container_width=True) and predio_quitar:
                        st.session_state["lista_predios_manual"].remove(predio_quitar)
                        st.rerun()

                if st.button("Limpiar toda la lista", use_container_width=False):
                    st.session_state["lista_predios_manual"] = []
                    st.rerun()

            predios_seleccionados = lista_actual
            st.session_state["predios_prog_manual"] = predios_seleccionados

        else:  # Carga masiva
            st.markdown('<div class="section-title">Carga masiva desde Excel</div>', unsafe_allow_html=True)
            archivo = st.file_uploader(
                "Suba un archivo Excel con columna REA",
                type=["xlsx", "xls"],
                help="El archivo debe tener al menos una columna llamada REA",
            )
            if archivo:
                try:
                    df_excel = pd.read_excel(archivo, dtype=str)
                    df_excel.columns = df_excel.columns.str.strip().str.upper()
                    if "REA" not in df_excel.columns:
                        st.error("El archivo no tiene columna REA")
                    else:
                        reas_excel = df_excel["REA"].dropna().str.strip().unique().tolist()
                        encontrados = maestro[maestro["REA"].isin(reas_excel)]
                        no_encontrados = [r for r in reas_excel if r not in maestro["REA"].values]

                        st.success(f"✓ {len(encontrados)} predios encontrados | ⚠️ {len(no_encontrados)} no encontrados")

                        if no_encontrados:
                            with st.expander(f"REAs no encontrados ({len(no_encontrados)})"):
                                st.write(no_encontrados)

                        show_cols = [c for c in ["REA", "CHIP", "LOCALIDAD", "BARRIO", "DIRECCION", "ESTADO_REA"] if c in encontrados.columns]
                        st.dataframe(encontrados[show_cols], use_container_width=True, hide_index=True)
                        predios_seleccionados = encontrados["REA"].tolist()
                        st.session_state["predios_prog_masiva"] = predios_seleccionados
                except Exception as e:
                    st.error(f"Error leyendo el archivo: {e}")

        # ── FORMULARIO DE PROGRAMACIÓN ─────────────────────────
        predios_a_programar = (
            predios_seleccionados
            or st.session_state.get("predios_prog_manual", [])
            or st.session_state.get("predios_prog_masiva", [])
        )

        if predios_a_programar:
            st.markdown('<div class="section-title">Datos de programación</div>', unsafe_allow_html=True)

            # ── Fecha / hora / observaciones (comunes) ─────────────────
            gcol1, gcol2 = st.columns(2)
            with gcol1:
                fecha_prog = st.date_input("Fecha programada", value=date.today(), key="fecha_prog_rea")
                hora_est   = st.time_input("Hora estimada", value=time(8, 0), key="hora_prog_rea")
            with gcol2:
                obs_prog = st.text_area("Observaciones de programación", height=100, key="obs_prog_rea")

            # ── Técnico por predio ─────────────────────────────────────
            st.markdown('<div class="section-title">Asignación de técnico por predio</div>', unsafe_allow_html=True)
            st.caption("Seleccione el técnico responsable de cada predio. Puede cambiar la asignación antes de confirmar.")

            # Inicializar asignaciones en session_state
            if "asig_tecnicos" not in st.session_state:
                st.session_state["asig_tecnicos"] = {}

            for rea in predios_a_programar:
                info = maestro[maestro["REA"] == rea]
                dir_label = info["DIRECCION"].iloc[0] if not info.empty and pd.notna(info["DIRECCION"].iloc[0]) else ""
                label = f"{rea}" + (f"  —  {dir_label}" if dir_label else "")
                prev = st.session_state["asig_tecnicos"].get(rea, [])
                tec_predio = st.multiselect(
                    label,
                    options=tecnicos_activos,
                    default=[t for t in prev if t in tecnicos_activos],
                    key=f"tec_predio_{rea}",
                )
                st.session_state["asig_tecnicos"][rea] = tec_predio

            if st.button("Programar Visita(s)", use_container_width=True, key="btn_prog_rea"):
                sin_tecnico = [r for r in predios_a_programar if not st.session_state["asig_tecnicos"].get(r)]
                if sin_tecnico:
                    st.warning(f"Faltan técnicos en: {', '.join(sin_tecnico)}")
                else:
                    visitas_df = load_visitas()
                    nuevas = []
                    nums_asignados = []
                    for rea in predios_a_programar:
                        tecs = st.session_state["asig_tecnicos"].get(rea, [])
                        num_visita = next_num_visita(visitas_df)
                        num_predio = count_visitas_predio(visitas_df, rea) + 1
                        nueva = {
                            "NUM_VISITA": num_visita,
                            "FECHA_PROGRAMADA": str(fecha_prog),
                            "REA": str(rea),
                            "SIN_REA": "No",
                            "DIRECCION_MANUAL": "",
                            "LATITUD_MANUAL": "",
                            "LONGITUD_MANUAL": "",
                            "TECNICOS": "|".join(tecs),
                            "ESTADO": "Pendiente",
                            "NUM_VISITA_PREDIO": str(num_predio),
                            "FECHA_REGISTRO": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "OBSERVACIONES_PROG": obs_prog,
                        }
                        nuevas.append(nueva)
                        nums_asignados.append(num_visita)
                        visitas_df = pd.concat([visitas_df, pd.DataFrame([nueva])], ignore_index=True)

                    save_visitas(visitas_df)
                    st.session_state["predios_prog_manual"] = []
                    st.session_state["predios_prog_masiva"] = []
                    st.session_state["lista_predios_manual"] = []
                    st.session_state["asig_tecnicos"] = {}
                    st.success(
                        f"✅ {len(nuevas)} visita(s) programada(s) exitosamente.\n\n"
                        f"Números asignados: {', '.join(nums_asignados)}"
                    )

    # ── TAB SIN REA ──────────────────────────────────────────────
    with tab_sin_rea:
        st.markdown('<div class="section-title">Predio sin REA (visita de campo)</div>', unsafe_allow_html=True)

        with st.form("form_sin_rea"):
            scol1, scol2 = st.columns(2)
            with scol1:
                dir_manual = st.text_input("Dirección", placeholder="Ej: CL 50 # 12-34")
                barrio_manual = st.text_input("Barrio")
                localidad_manual = st.text_input("Localidad")
            with scol2:
                lat_manual = st.text_input("Latitud (opcional)", placeholder="4.6xxx")
                lon_manual = st.text_input("Longitud (opcional)", placeholder="-74.0xxx")
                if lat_manual and lon_manual:
                    st.markdown(
                        f"🗺️ [Ver en Maps](https://www.google.com/maps?q={lat_manual},{lon_manual})"
                    )

            st.markdown("---")
            sfcol1, sfcol2 = st.columns(2)
            with sfcol1:
                fecha_sin = st.date_input("Fecha programada", value=date.today(), key="fecha_sin")
                tecnicos_sin = st.multiselect(
                    "Técnicos asignados",
                    options=tecnicos_activos,
                    key="tec_sin",
                )
            with sfcol2:
                hora_sin = st.time_input("Hora estimada", value=time(8, 0), key="hora_sin")
                obs_sin = st.text_area("Observaciones", height=80, key="obs_sin")

            sub_sin = st.form_submit_button("Programar Visita Sin REA", use_container_width=True)

        if sub_sin:
            if not dir_manual:
                st.warning("Ingrese al menos una dirección.")
            elif not tecnicos_sin:
                st.warning("Seleccione al menos un técnico.")
            else:
                visitas_df = load_visitas()
                num_visita = next_num_visita(visitas_df)
                nueva = {
                    "NUM_VISITA": num_visita,
                    "FECHA_PROGRAMADA": str(fecha_sin),
                    "REA": "",
                    "SIN_REA": "Si",
                    "DIRECCION_MANUAL": f"{dir_manual}, {barrio_manual}, {localidad_manual}".strip(", "),
                    "LATITUD_MANUAL": lat_manual,
                    "LONGITUD_MANUAL": lon_manual,
                    "TECNICOS": "|".join(tecnicos_sin),
                    "ESTADO": "Pendiente",
                    "NUM_VISITA_PREDIO": "1",
                    "FECHA_REGISTRO": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "OBSERVACIONES_PROG": obs_sin,
                }
                visitas_df = pd.concat([visitas_df, pd.DataFrame([nueva])], ignore_index=True)
                save_visitas(visitas_df)
                st.success(f"✅ Visita programada. Número asignado: **{num_visita}**")


# ════════════════════════════════════════════════════════════════
# PAGE: VISITAS PROGRAMADAS
# ════════════════════════════════════════════════════════════════
elif pagina == "Visitas Programadas":
    st.markdown(
        '<div class="cvp-header"><h1>Visitas Programadas</h1><p>Predios pendientes de visita técnica</p></div>',
        unsafe_allow_html=True,
    )

    visitas = load_visitas()
    maestro = load_maestro()

    if visitas.empty:
        st.info("No hay visitas programadas aún. Vaya a 📋 Programar Visita para crear la primera.")
    else:
        # Solo pendientes por defecto
        pendientes_vp = visitas[visitas["ESTADO"] == "Pendiente"].copy() if "ESTADO" in visitas.columns else visitas.copy()

        # Merge con maestro para localidad, dirección, coordenadas
        if not maestro.empty and "REA" in pendientes_vp.columns:
            merge_cols = [c for c in ["REA", "LOCALIDAD", "BARRIO", "DIRECCION", "ESTADO_REA",
                                       "ENLACE_MAPS", "LATITUD", "LONGITUD"] if c in maestro.columns]
            vm = pendientes_vp.merge(
                maestro[merge_cols].rename(columns={
                    "LOCALIDAD": "LOCALIDAD_M", "BARRIO": "BARRIO_M",
                    "DIRECCION": "DIRECCION_M", "LATITUD": "LAT_M", "LONGITUD": "LON_M",
                }),
                on="REA", how="left",
            )
        else:
            vm = pendientes_vp.copy()

        # ── FILTROS ───────────────────────────────────────────
        st.markdown('<div class="section-title">Filtros</div>', unsafe_allow_html=True)
        fcol1, fcol2, fcol3 = st.columns(3)
        with fcol1:
            tec_options = ["Todos"]
            if "TECNICOS" in vm.columns:
                all_tecs = set()
                for t in vm["TECNICOS"].dropna():
                    all_tecs.update(str(t).split("|"))
                tec_options += sorted(all_tecs)
            tec_filtro = st.selectbox("Técnico", tec_options, key="vp_tec")
        with fcol2:
            loc_options = ["Todas"]
            loc_col = "LOCALIDAD_M" if "LOCALIDAD_M" in vm.columns else None
            if loc_col:
                loc_options += sorted(vm[loc_col].dropna().unique().tolist())
            loc_filtro = st.selectbox("Localidad", loc_options, key="vp_loc")
        with fcol3:
            try:
                f_min = date.fromisoformat(str(vm["FECHA_PROGRAMADA"].min())[:10])
            except Exception:
                f_min = date.today()
            fechas_vp = st.date_input("Rango fechas", value=(f_min, date.today()), key="vp_fechas")

        df_f = vm.copy()
        if tec_filtro != "Todos" and "TECNICOS" in df_f.columns:
            df_f = df_f[df_f["TECNICOS"].str.contains(tec_filtro, na=False)]
        if loc_filtro != "Todas" and loc_col and loc_col in df_f.columns:
            df_f = df_f[df_f[loc_col] == loc_filtro]
        if isinstance(fechas_vp, (list, tuple)) and len(fechas_vp) == 2 and "FECHA_PROGRAMADA" in df_f.columns:
            df_f["_fd"] = pd.to_datetime(df_f["FECHA_PROGRAMADA"], errors="coerce")
            df_f = df_f[(df_f["_fd"] >= pd.Timestamp(fechas_vp[0])) & (df_f["_fd"] <= pd.Timestamp(fechas_vp[1]))]
            df_f = df_f.drop(columns=["_fd"])

        st.caption(f"**{len(df_f)} visitas pendientes** con los filtros actuales.")

        # ── TABLA ─────────────────────────────────────────────
        tab_lista, tab_mapa, tab_eliminar = st.tabs(["📋 Lista", "🗺️ Mapa", "🗑️ Eliminar predio"])

        with tab_lista:
            show_cols = [c for c in [
                "NUM_VISITA", "FECHA_PROGRAMADA", "REA", "TECNICOS",
                "LOCALIDAD_M", "BARRIO_M", "DIRECCION_M", "ESTADO_REA", "ENLACE_MAPS",
            ] if c in df_f.columns]
            st.dataframe(
                df_f[show_cols],
                use_container_width=True,
                hide_index=True,
                column_config={
                    "ENLACE_MAPS": st.column_config.LinkColumn("Maps", display_text="📍 Ver mapa"),
                    "NUM_VISITA": st.column_config.TextColumn("N° Visita"),
                    "FECHA_PROGRAMADA": st.column_config.TextColumn("Fecha"),
                    "TECNICOS": st.column_config.TextColumn("Técnico(s)"),
                    "LOCALIDAD_M": st.column_config.TextColumn("Localidad"),
                    "BARRIO_M": st.column_config.TextColumn("Barrio"),
                    "DIRECCION_M": st.column_config.TextColumn("Dirección"),
                    "ESTADO_REA": st.column_config.TextColumn("Estado REA"),
                },
            )

            # Exportar Excel con columnas limpias
            export_cols = [c for c in [
                "NUM_VISITA", "FECHA_PROGRAMADA", "REA", "TECNICOS",
                "LOCALIDAD_M", "BARRIO_M", "DIRECCION_M", "ESTADO_REA",
                "LAT_M", "LON_M", "ENLACE_MAPS",
            ] if c in df_f.columns]
            df_export = df_f[export_cols].rename(columns={
                "LOCALIDAD_M": "LOCALIDAD", "BARRIO_M": "BARRIO",
                "DIRECCION_M": "DIRECCION", "LAT_M": "LATITUD", "LON_M": "LONGITUD",
            })
            buf_vp = io.BytesIO()
            with pd.ExcelWriter(buf_vp, engine="openpyxl") as writer:
                df_export.to_excel(writer, index=False, sheet_name="Pendientes")
            st.download_button(
                "⬇️ Exportar a Excel",
                data=buf_vp.getvalue(),
                file_name=f"visitas_pendientes_{date.today()}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        # ── MAPA ──────────────────────────────────────────────
        with tab_mapa:
            # Obtener coordenadas: preferir maestro, fallback a manual
            df_map = df_f.copy()
            df_map["_lat"] = pd.to_numeric(df_map.get("LAT_M"), errors="coerce")
            df_map["_lon"] = pd.to_numeric(df_map.get("LON_M"), errors="coerce")
            # Para sin REA, usar coordenadas manuales
            if "LATITUD_MANUAL" in df_map.columns:
                df_map["_lat"] = df_map["_lat"].fillna(pd.to_numeric(df_map["LATITUD_MANUAL"], errors="coerce"))
                df_map["_lon"] = df_map["_lon"].fillna(pd.to_numeric(df_map["LONGITUD_MANUAL"], errors="coerce"))

            df_map_clean = df_map.dropna(subset=["_lat", "_lon"])

            if df_map_clean.empty:
                st.info("No hay coordenadas disponibles para los predios filtrados. Verifique que el maestro tenga LATITUD y LONGITUD.")
            else:
                n_sin_coord = len(df_map) - len(df_map_clean)
                if n_sin_coord > 0:
                    st.caption(f"⚠️ {n_sin_coord} predio(s) sin coordenadas no aparecen en el mapa.")

                # Mapa con pydeck (scatter layer)
                try:
                    import pydeck as pdk
                    dir_col = df_map_clean.get("DIRECCION_M", pd.Series("", index=df_map_clean.index)) if "DIRECCION_M" in df_map_clean.columns else pd.Series("", index=df_map_clean.index)
                    tec_col = df_map_clean["TECNICOS"] if "TECNICOS" in df_map_clean.columns else pd.Series("", index=df_map_clean.index)
                    df_map_clean = df_map_clean.copy()
                    df_map_clean["tooltip"] = (
                        df_map_clean["REA"].fillna("") + "\n" +
                        dir_col.fillna("") + "\n" +
                        tec_col.fillna("")
                    )
                    layer = pdk.Layer(
                        "ScatterplotLayer",
                        data=df_map_clean[["_lat", "_lon", "tooltip"]].rename(columns={"_lat": "lat", "_lon": "lon"}),
                        get_position="[lon, lat]",
                        get_color=[204, 0, 0, 210],
                        get_radius=55,
                        radius_min_pixels=6,
                        radius_max_pixels=18,
                        pickable=True,
                    )
                    view = pdk.ViewState(
                        latitude=df_map_clean["_lat"].mean(),
                        longitude=df_map_clean["_lon"].mean(),
                        zoom=12,
                        pitch=0,
                    )
                    st.pydeck_chart(pdk.Deck(
                        layers=[layer],
                        initial_view_state=view,
                        tooltip={"text": "{tooltip}"},
                        map_style="https://basemaps.cartocdn.com/gl/positron-gl-style/style.json",
                    ))
                except ImportError:
                    st.map(
                        df_map_clean.rename(columns={"_lat": "latitude", "_lon": "longitude"})[["latitude", "longitude"]],
                        zoom=11,
                    )
                    st.caption("Instale `pydeck` para mapa interactivo con tooltips: `pip install pydeck`")

        # ── ELIMINAR PREDIO ───────────────────────────────────
        with tab_eliminar:
            st.markdown('<div class="section-title">Eliminar predio de una visita programada</div>', unsafe_allow_html=True)
            st.warning("⚠️ Esta acción elimina permanentemente el registro de programación. Solo aplica a visitas **Pendientes**.")

            if pendientes_vp.empty:
                st.info("No hay visitas pendientes para eliminar.")
            else:
                opciones_elim = {}
                for _, row in pendientes_vp.iterrows():
                    nv = row.get("NUM_VISITA", "")
                    rea = row.get("REA", "") or row.get("DIRECCION_MANUAL", "Sin REA")
                    lbl = f"{nv} — {rea}"
                    opciones_elim[lbl] = nv

                sel_elim = st.selectbox("Seleccione la visita a eliminar", list(opciones_elim.keys()), key="sel_elim")
                nv_elim = opciones_elim[sel_elim]

                st.info(f"Va a eliminar: **{sel_elim}**")

                confirmar = st.checkbox("Confirmo que deseo eliminar esta visita programada", key="confirm_elim")
                if st.button("🗑️ Eliminar visita", disabled=not confirmar, key="btn_elim"):
                    delete_visita_supabase(nv_elim)
                    visitas_df = load_visitas()
                    visitas_df = visitas_df[visitas_df["NUM_VISITA"] != nv_elim]
                    save_visitas(visitas_df)
                    st.success(f"✅ Visita **{nv_elim}** eliminada.")
                    st.rerun()


# ════════════════════════════════════════════════════════════════
# PAGE: REGISTRAR RESULTADO
# ════════════════════════════════════════════════════════════════
elif pagina == "Registrar Resultado":
    st.markdown(
        '<div class="cvp-header"><h1>Registrar Resultado de Visita</h1><p>Documente el resultado de una visita técnica programada</p></div>',
        unsafe_allow_html=True,
    )

    visitas = load_visitas()
    maestro = load_maestro()
    resultados = load_resultados()
    tecnicos_df = load_tecnicos()
    tecnicos_activos = (
        tecnicos_df[tecnicos_df["ACTIVO"] == True]["NOMBRE"].tolist()
        if not tecnicos_df.empty else []
    )

    # ── DESCARGA INLINE ─────────────────────────────────────────
    descarga_nv = st.session_state.get("mostrar_descarga")
    if descarga_nv:
        res_dl_df = load_resultados()
        if not res_dl_df.empty and descarga_nv in res_dl_df["NUM_VISITA"].values:
            res_dl = res_dl_df[res_dl_df["NUM_VISITA"] == descarga_nv].iloc[0]
            vis_dl_match = visitas[visitas["NUM_VISITA"] == descarga_nv]
            vis_dl = vis_dl_match.iloc[0] if not vis_dl_match.empty else pd.Series()
            rea_dl = _str(res_dl.get("REA", ""))
            maestro_dl = pd.Series()
            if rea_dl and not maestro.empty:
                m_dl = maestro[maestro["REA"] == rea_dl]
                if not m_dl.empty:
                    maestro_dl = m_dl.iloc[0]
            resultado_dl = _str(res_dl.get("RESULTADO", ""))

            st.success(f"✅ Resultado **{resultado_dl}** guardado para visita **{descarga_nv}**")
            st.markdown('<div class="section-title">📄 Descargar Formato Oficial</div>', unsafe_allow_html=True)

            if resultado_dl == "Exitosa":
                st.markdown('<div class="banner-exitosa">VISITA EXITOSA — Ficha Técnica de Reconocimiento (208-REAS-Ft-30)</div>', unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)
                if not FICHA_TEMPLATE.exists():
                    st.error(f"Plantilla no encontrada: {FICHA_TEMPLATE}")
                else:
                    xlsx_bytes = generar_ficha_tecnica(res_dl, vis_dl, maestro_dl)
                    st.download_button(
                        label="⬇️ Descargar Ficha Técnica XLSX",
                        data=xlsx_bytes,
                        file_name=f"FichaTecnica_{descarga_nv}_{rea_dl}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                    )
                    st.caption("Archivo XLSX pre-diligenciado. Complete los campos restantes en Excel.")
            else:
                st.markdown('<div class="banner-fallida">VISITA FALLIDA — Informe de Gestión (208-REAS-Ft-176)</div>', unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)
                if not INFORME_TEMPLATE.exists():
                    st.error(f"Plantilla no encontrada: {INFORME_TEMPLATE}")
                else:
                    docx_bytes = generar_informe_gestion(res_dl, vis_dl, maestro_dl)
                    st.download_button(
                        label="⬇️ Descargar Informe de Gestión DOCX",
                        data=docx_bytes,
                        file_name=f"InformeGestion_{descarga_nv}_{rea_dl}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                    st.caption("Archivo DOCX pre-diligenciado. Complete los campos restantes en Word.")

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("📋 Registrar otra visita", use_container_width=False):
                st.session_state["mostrar_descarga"] = None
                st.rerun()

    else:
        # ── FORMULARIO DE REGISTRO ───────────────────────────────
        if visitas.empty:
            st.info("No hay visitas programadas. Programe una visita primero.")
        else:
            pendientes = visitas[visitas["ESTADO"] == "Pendiente"] if "ESTADO" in visitas.columns else visitas

            if pendientes.empty:
                st.info("No hay visitas pendientes de registro.")
            else:
                # Build dropdown options
                def visita_label(row):
                    rea = row.get("REA", "")
                    dir_m = row.get("DIRECCION_MANUAL", "")
                    nv = row.get("NUM_VISITA", "")
                    desc = rea if rea else dir_m
                    return f"{nv} | {desc}"

                opciones = {visita_label(row): row for _, row in pendientes.iterrows()}

                presel = st.session_state.get("registro_num_visita")
                presel_key = None
                if presel:
                    for k, v in opciones.items():
                        if v.get("NUM_VISITA") == presel:
                            presel_key = k
                            break

                sel_label = st.selectbox(
                    "Seleccione la visita a registrar",
                    list(opciones.keys()),
                    index=list(opciones.keys()).index(presel_key) if presel_key else 0,
                )
                visita_sel = opciones[sel_label]
                num_visita_sel = visita_sel.get("NUM_VISITA", "")
                rea_sel = visita_sel.get("REA", "")
                num_predio = visita_sel.get("NUM_VISITA_PREDIO", "1")

                # Pre-cargar UPL/UPZ desde maestro si están disponibles
                _upl_default = ""
                _upz_default = ""
                if rea_sel and not maestro.empty:
                    _pi_pre = maestro[maestro["REA"] == rea_sel]
                    if not _pi_pre.empty:
                        _upl_default = _str(_pi_pre.iloc[0].get("UPL", ""))
                        _upz_default = _str(_pi_pre.iloc[0].get("UPZ", ""))
                # Inyectar en session_state ANTES del form para que value= funcione
                _upl_key = f"upl_e_{rea_sel}"
                _upz_key = f"upz_e_{rea_sel}"
                if st.session_state.get("_last_rea_form") != rea_sel:
                    st.session_state[_upl_key] = _upl_default
                    st.session_state[_upz_key] = _upz_default
                    st.session_state["_last_rea_form"] = rea_sel

                # Predio info card
                if rea_sel and not maestro.empty:
                    predio_info = maestro[maestro["REA"] == rea_sel]
                    if not predio_info.empty:
                        pi = predio_info.iloc[0]
                        st.markdown('<div class="section-title">Información del Predio</div>', unsafe_allow_html=True)
                        ic1, ic2, ic3 = st.columns(3)
                        with ic1:
                            st.markdown(f"**REA:** {pi.get('REA','')}")
                            st.markdown(f"**CHIP:** {pi.get('CHIP','')}")
                            st.markdown(f"**Localidad:** {pi.get('LOCALIDAD','')}")
                        with ic2:
                            st.markdown(f"**Barrio:** {pi.get('BARRIO','')}")
                            st.markdown(f"**Dirección:** {pi.get('DIRECCION','')}")
                            st.markdown(f"**Estado REA:** {pi.get('ESTADO_REA','')}")
                        with ic3:
                            st.markdown(f"**Propietario:** {pi.get('PROPIETARIO_1','')}")
                            st.markdown(f"**Cédula:** {pi.get('CEDULA_1','')}")
                            avaluo = pi.get('AVALUO', 0)
                            try:
                                avaluo_fmt = f"${float(avaluo):,.0f}" if float(str(avaluo)) > 0 else "No registrado"
                            except Exception:
                                avaluo_fmt = str(avaluo)
                            st.markdown(f"**Avalúo:** {avaluo_fmt}")

                st.info(f"Esta es la **visita N° {num_predio}** a este predio.")

                # ── RESULTADO FORM ────────────────────────────────
                st.markdown('<div class="section-title">Resultado de la Visita</div>', unsafe_allow_html=True)

                # Radio FUERA del form para que cambie el formulario dinámicamente
                resultado = st.radio("Resultado", ["Exitosa", "Fallida"], horizontal=True, key="resultado_radio")

                with st.form("form_resultado"):
                    fecha_visita = st.date_input("Fecha de visita", value=date.today())

                    # Técnicos programados como default, editables
                    tecs_programados = [t.strip() for t in str(visita_sel.get("TECNICOS", "")).split("|") if t.strip()]
                    tecs_validos = [t for t in tecs_programados if t in tecnicos_activos]
                    tecs_resultado = st.multiselect(
                        "Técnico(s) que realizaron la visita",
                        options=tecnicos_activos,
                        default=tecs_validos,
                        help="Pre-cargado con los técnicos programados. Modifique si hubo cambios.",
                    )
                    tec_display = "|".join(tecs_resultado)

                    if resultado == "Exitosa":
                        rc1, rc2, rc3 = st.columns(3)
                        with rc1:
                            hora_ini = st.time_input("Hora inicio", value=time(9, 0))
                            ocupacion = st.selectbox("Ocupación", OCUPACIONES)
                            tipo_const = st.selectbox("Tipo de construcción", TIPOS_CONSTRUCCION)
                            num_pisos = st.selectbox("Número de pisos", PISOS_OPTIONS)
                        with rc2:
                            hora_fin = st.time_input("Hora fin", value=time(10, 0))
                            prop_contactado = st.radio("Propietario contactado", ["Si", "No"], horizontal=True)
                            estado_cons = st.selectbox("Estado de conservación", ESTADOS_CONSERVACION)
                        with rc3:
                            tipo_inmueble = st.selectbox("Tipo de inmueble", ["Casa", "Casa Lote", "Institucional", "Industrial", "Bodega", "Garaje", "Oficina", "Lote"])
                            estrato = st.selectbox("Estrato", ["1", "2", "3", "4", "5", "6"])
                            unidades_vivienda = st.number_input("Unidades de vivienda", min_value=1, max_value=20, value=1, step=1)

                        ra1, ra2, ra3, ra4 = st.columns(4)
                        with ra1:
                            upl = st.text_input("UPL", key=_upl_key)
                        with ra2:
                            upz = st.text_input("UPZ", key=_upz_key)
                        with ra3:
                            area_terreno = st.text_input("Área terreno (m²)", key="at_e")
                        with ra4:
                            area_construccion = st.text_input("Área construcción (m²)", key="ac_e")

                        st.markdown("**Linderos**")
                        lc1, lc2, lc3, lc4 = st.columns(4)
                        with lc1:
                            lindero_n = st.text_input("Norte", key="ln")
                        with lc2:
                            lindero_s = st.text_input("Sur", key="ls")
                        with lc3:
                            lindero_o = st.text_input("Oriente", key="lor")
                        with lc4:
                            lindero_oc = st.text_input("Occidente", key="loc")

                        observaciones = st.text_area("Observaciones", height=100)

                        motivo_fallida = ""
                        hora_ini_f = str(hora_ini)
                        hora_fin_f = str(hora_fin)
                        tipo_gestion = ""
                        telefono_beneficiario = ""
                        correo_beneficiario = ""
                        componente = ""

                    else:  # Fallida
                        fb1, fb2, fb3 = st.columns(3)
                        with fb1:
                            hora_ini_fall = st.time_input("Hora inicio", value=time(9, 0), key="hi_fall")
                            tipo_gestion = st.selectbox(
                                "Tipo de gestión",
                                ["Visita a campo", "Seguimiento al proceso", "Entrega de documentos"],
                                key="tg_fall",
                            )
                        with fb2:
                            hora_fin_fall = st.time_input("Hora fin", value=time(10, 0), key="hf_fall")
                            componente = st.text_input("Componente", value="Reasentamiento", key="comp_fall")
                        with fb3:
                            telefono_beneficiario = st.text_input("Teléfono beneficiario", key="tel_fall")
                            correo_beneficiario = st.text_input("Correo beneficiario", key="email_fall")

                        st.markdown("**Motivo de visita fallida**")
                        motivos_sel = []
                        mf_cols = st.columns(2)
                        for i, mot in enumerate(MOTIVOS_FALLIDA):
                            with mf_cols[i % 2]:
                                if st.checkbox(mot, key=f"mot_{i}"):
                                    motivos_sel.append(mot)

                        observaciones = st.text_area("Descripción de la gestión realizada / Observaciones", height=120, key="obs_fallida")

                        motivo_fallida = "|".join(motivos_sel)
                        hora_ini_f = str(hora_ini_fall)
                        hora_fin_f = str(hora_fin_fall)
                        ocupacion = prop_contactado = tipo_const = num_pisos = estado_cons = ""
                        tipo_inmueble = estrato = ""
                        unidades_vivienda = 0
                        lindero_n = lindero_s = lindero_o = lindero_oc = ""
                        upl = upz = area_terreno = area_construccion = ""

                    sub_res = st.form_submit_button("Guardar Resultado", use_container_width=True)

                if sub_res:
                    # Build result row
                    nuevo_resultado = {
                        "NUM_VISITA": num_visita_sel,
                        "REA": rea_sel,
                        "FECHA_VISITA": str(fecha_visita),
                        "HORA_INICIO": hora_ini_f if resultado == "Exitosa" else "",
                        "HORA_FIN": hora_fin_f if resultado == "Exitosa" else "",
                        "TECNICOS": tec_display,
                        "RESULTADO": resultado,
                        "OCUPACION": ocupacion if resultado == "Exitosa" else "",
                        "PROP_CONTACTADO": prop_contactado if resultado == "Exitosa" else "",
                        "TIPO_CONSTRUCCION": tipo_const if resultado == "Exitosa" else "",
                        "NUM_PISOS": num_pisos if resultado == "Exitosa" else "",
                        "ESTADO_CONSERVACION": estado_cons if resultado == "Exitosa" else "",
                        "LINDERO_NORTE": lindero_n if resultado == "Exitosa" else "",
                        "LINDERO_SUR": lindero_s if resultado == "Exitosa" else "",
                        "LINDERO_ORIENTE": lindero_o if resultado == "Exitosa" else "",
                        "LINDERO_OCCIDENTE": lindero_oc if resultado == "Exitosa" else "",
                        "TIPO_INMUEBLE": tipo_inmueble if resultado == "Exitosa" else "",
                        "ESTRATO": str(estrato) if resultado == "Exitosa" else "",
                        "UNIDADES_VIVIENDA": str(int(unidades_vivienda)) if resultado == "Exitosa" else "",
                        "UPL": upl if resultado == "Exitosa" else "",
                        "UPZ": upz if resultado == "Exitosa" else "",
                        "AREA_TERRENO": area_terreno if resultado == "Exitosa" else "",
                        "AREA_CONSTRUCCION": area_construccion if resultado == "Exitosa" else "",
                        "TIPO_GESTION": tipo_gestion if resultado == "Fallida" else "",
                        "TELEFONO_BENEFICIARIO": telefono_beneficiario if resultado == "Fallida" else "",
                        "CORREO_BENEFICIARIO": correo_beneficiario if resultado == "Fallida" else "",
                        "COMPONENTE": componente if resultado == "Fallida" else "",
                        "MOTIVO_FALLIDA": motivo_fallida if resultado == "Fallida" else "",
                        "OBSERVACIONES": observaciones,
                        "FOTOS": "",
                        "FECHA_REGISTRO": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    }

                    resultados_df = load_resultados()
                    resultados_df = pd.concat(
                        [resultados_df, pd.DataFrame([nuevo_resultado])], ignore_index=True
                    )
                    save_resultados(resultados_df)

                    # Update visitas estado
                    visitas_df = load_visitas()
                    visitas_df.loc[visitas_df["NUM_VISITA"] == num_visita_sel, "ESTADO"] = resultado
                    save_visitas(visitas_df)

                    st.session_state["pdf_num_visita"] = num_visita_sel
                    st.session_state["registro_num_visita"] = None
                    st.session_state["mostrar_descarga"] = num_visita_sel
                    st.rerun()


# ════════════════════════════════════════════════════════════════
# PAGE: GESTIÓN DE TÉCNICOS
# ════════════════════════════════════════════════════════════════
elif pagina == "Gestión de Técnicos":
    st.markdown(
        '<div class="cvp-header"><h1>Gestión de Técnicos</h1><p>Administre el equipo técnico de reasentamientos</p></div>',
        unsafe_allow_html=True,
    )

    tecnicos_df = load_tecnicos()

    st.markdown('<div class="section-title">Técnicos registrados</div>', unsafe_allow_html=True)

    if not tecnicos_df.empty:
        edited = st.data_editor(
            tecnicos_df,
            use_container_width=True,
            num_rows="fixed",
            column_config={
                "ID_TECNICO": st.column_config.TextColumn("ID", disabled=True),
                "NOMBRE": st.column_config.TextColumn("Nombre completo"),
                "CARGO": st.column_config.SelectboxColumn("Cargo", options=CARGOS),
                "EMAIL": st.column_config.TextColumn("Correo electrónico"),
                "CONTRATO": st.column_config.TextColumn("N° Contrato"),
                "ACTIVO": st.column_config.CheckboxColumn("Activo"),
            },
            hide_index=True,
        )

        if st.button("💾 Guardar cambios", key="save_tec_edit"):
            save_tecnicos(edited)
            st.success("✅ Cambios guardados exitosamente.")
            st.rerun()

    st.markdown('<div class="section-title">Agregar nuevo técnico</div>', unsafe_allow_html=True)

    with st.form("form_nuevo_tecnico"):
        nc1, nc2 = st.columns(2)
        with nc1:
            nuevo_nombre = st.text_input("Nombre completo")
            nuevo_cargo = st.selectbox("Cargo", CARGOS)
            nuevo_contrato = st.text_input("N° Contrato")
        with nc2:
            nuevo_email = st.text_input("Correo electrónico")
            nuevo_activo = st.checkbox("Activo", value=True)

        sub_tec = st.form_submit_button("Agregar Técnico", use_container_width=True)

    if sub_tec:
        if not nuevo_nombre:
            st.warning("Ingrese el nombre del técnico.")
        else:
            tecnicos_df = load_tecnicos()
            # Generate new ID
            if tecnicos_df.empty:
                new_id = "T001"
            else:
                ids = tecnicos_df["ID_TECNICO"].str.replace("T", "", regex=False)
                max_id = ids.apply(lambda x: int(x) if x.isdigit() else 0).max()
                new_id = f"T{str(max_id + 1).zfill(3)}"

            nuevo_tecnico = {
                "ID_TECNICO": new_id,
                "NOMBRE": nuevo_nombre,
                "CARGO": nuevo_cargo,
                "EMAIL": nuevo_email,
                "CONTRATO": nuevo_contrato,
                "ACTIVO": nuevo_activo,
            }
            tecnicos_df = pd.concat(
                [tecnicos_df, pd.DataFrame([nuevo_tecnico])], ignore_index=True
            )
            save_tecnicos(tecnicos_df)
            st.success(f"✅ Técnico '{nuevo_nombre}' agregado con ID {new_id}.")
            st.rerun()


# ════════════════════════════════════════════════════════════════
# PAGE: INDICADORES
# ════════════════════════════════════════════════════════════════
elif pagina == "Indicadores":
    st.markdown(
        '<div class="cvp-header"><h1>Indicadores de Gestión</h1><p>Análisis y estadísticas del programa de visitas técnicas</p></div>',
        unsafe_allow_html=True,
    )

    visitas = load_visitas()
    resultados = load_resultados()
    tecnicos_df = load_tecnicos()

    tab_gen, tab_tec, tab_pred, tab_rep = st.tabs(["📊 General", "👤 Por Técnico", "🏘️ Por Predio", "📥 Reporte Excel"])

    # ── TAB GENERAL ───────────────────────────────────────────
    with tab_gen:
        if visitas.empty:
            st.info("No hay datos de visitas aún.")
        else:
            # Date filter
            gf1, gf2 = st.columns([2, 2])
            with gf1:
                periodo = st.selectbox("Período", ["Esta semana", "Este mes", "Rango personalizado"])
            with gf2:
                if periodo == "Rango personalizado":
                    rango_g = st.date_input("Rango", value=(date.today().replace(day=1), date.today()), key="rango_g")
                else:
                    rango_g = None

            vis_g = visitas.copy()
            if "FECHA_PROGRAMADA" in vis_g.columns:
                vis_g["_fecha"] = pd.to_datetime(vis_g["FECHA_PROGRAMADA"], errors="coerce")
                today = pd.Timestamp(date.today())
                if periodo == "Esta semana":
                    start = today - pd.Timedelta(days=today.dayofweek)
                    vis_g = vis_g[vis_g["_fecha"] >= start]
                elif periodo == "Este mes":
                    vis_g = vis_g[vis_g["_fecha"].dt.month == today.month]
                elif rango_g and isinstance(rango_g, (list, tuple)) and len(rango_g) == 2:
                    vis_g = vis_g[
                        (vis_g["_fecha"] >= pd.Timestamp(rango_g[0]))
                        & (vis_g["_fecha"] <= pd.Timestamp(rango_g[1]))
                    ]

            tot_g = len(vis_g)
            exit_g = int((vis_g["ESTADO"] == "Exitosa").sum()) if "ESTADO" in vis_g.columns else 0
            fall_g = int((vis_g["ESTADO"] == "Fallida").sum()) if "ESTADO" in vis_g.columns else 0
            pct_g = round(exit_g / tot_g * 100, 1) if tot_g > 0 else 0

            mc1, mc2, mc3, mc4 = st.columns(4)
            mc1.metric("Total visitas", tot_g)
            mc2.metric("Exitosas", exit_g)
            mc3.metric("Fallidas", fall_g)
            mc4.metric("% Éxito", f"{pct_g}%")

            if not vis_g.empty and "ESTADO" in vis_g.columns and "_fecha" in vis_g.columns:
                vis_g["_semana"] = vis_g["_fecha"].dt.to_period("W").apply(
                    lambda r: str(r.start_time.date()) if pd.notna(r) else None
                )
                sem_data = vis_g.groupby(["_semana", "ESTADO"]).size().reset_index(name="count")

                if not sem_data.empty:
                    gc1, gc2 = st.columns([3, 2])
                    with gc1:
                        fig_bar = go.Figure()
                        for estado, color in [("Pendiente", COLOR_WARNING), ("Fallida", COLOR_DANGER), ("Exitosa", COLOR_SUCCESS)]:
                            d = sem_data[sem_data["ESTADO"] == estado]
                            if not d.empty:
                                fig_bar.add_trace(go.Bar(
                                    x=d["_semana"], y=d["count"], name=estado,
                                    marker_color=color, marker_line_width=0,
                                    hovertemplate="<b>%{x}</b><br>" + estado + ": %{y}<extra></extra>",
                                ))
                        fig_bar.update_layout(
                            barmode="stack",
                            title=dict(text="Visitas por semana", font=dict(size=15, color="#333"), x=0.01),
                            height=340,
                            margin=dict(l=10, r=10, t=50, b=10),
                            plot_bgcolor="white", paper_bgcolor="white",
                            legend=dict(orientation="h", yanchor="bottom", y=1.01, xanchor="right", x=1, font=dict(size=11)),
                            xaxis=dict(showgrid=False, tickfont=dict(size=10), title=""),
                            yaxis=dict(showgrid=True, gridcolor="#f5f5f5", tickfont=dict(size=10), title="Visitas"),
                        )
                        st.plotly_chart(fig_bar, use_container_width=True)

                    with gc2:
                        pie_data = vis_g["ESTADO"].value_counts().reset_index()
                        pie_data.columns = ["ESTADO", "count"]
                        color_map = {"Exitosa": COLOR_SUCCESS, "Fallida": COLOR_DANGER, "Pendiente": COLOR_WARNING}
                        fig_pie = go.Figure(go.Pie(
                            labels=pie_data["ESTADO"],
                            values=pie_data["count"],
                            hole=0.52,
                            marker=dict(colors=[color_map.get(e, "#999") for e in pie_data["ESTADO"]],
                                        line=dict(color="white", width=2)),
                            textinfo="label+percent",
                            textfont=dict(size=12),
                            hovertemplate="<b>%{label}</b><br>%{value} visitas (%{percent})<extra></extra>",
                        ))
                        fig_pie.update_layout(
                            title=dict(text="Distribución por resultado", font=dict(size=15, color="#333"), x=0.01),
                            height=340,
                            margin=dict(l=10, r=10, t=50, b=10),
                            paper_bgcolor="white",
                            showlegend=True,
                            legend=dict(orientation="h", yanchor="bottom", y=-0.15, xanchor="center", x=0.5, font=dict(size=11)),
                            annotations=[dict(
                                text=f"<b>{tot_g}</b><br>total",
                                x=0.5, y=0.5, font_size=14, showarrow=False,
                                font=dict(color="#333"),
                            )],
                        )
                        st.plotly_chart(fig_pie, use_container_width=True)

    # ── TAB POR TÉCNICO ───────────────────────────────────────
    with tab_tec:
        if visitas.empty or "TECNICOS" not in visitas.columns:
            st.info("No hay datos suficientes.")
        else:
            # Explode technicians
            vis_tec = visitas.copy()
            vis_tec["_fecha"] = pd.to_datetime(vis_tec.get("FECHA_PROGRAMADA"), errors="coerce")
            vis_tec_exp = vis_tec.assign(
                TECNICO=vis_tec["TECNICOS"].str.split("|")
            ).explode("TECNICO")
            vis_tec_exp["TECNICO"] = vis_tec_exp["TECNICO"].str.strip()

            resumen_tec = vis_tec_exp.groupby("TECNICO").agg(
                Total=("NUM_VISITA", "count"),
                Exitosas=("ESTADO", lambda x: (x == "Exitosa").sum()),
                Fallidas=("ESTADO", lambda x: (x == "Fallida").sum()),
            ).reset_index()
            resumen_tec["% Éxito"] = (resumen_tec["Exitosas"] / resumen_tec["Total"] * 100).round(1)

            st.markdown('<div class="section-title">Resumen por técnico</div>', unsafe_allow_html=True)
            st.dataframe(resumen_tec, use_container_width=True, hide_index=True)

            resumen_sorted = resumen_tec.sort_values("Total", ascending=True)
            fig_tec = go.Figure()
            fig_tec.add_trace(go.Bar(
                y=resumen_sorted["TECNICO"], x=resumen_sorted["Exitosas"],
                name="Exitosas", orientation="h",
                marker=dict(color=COLOR_SUCCESS, line_width=0),
                hovertemplate="<b>%{y}</b><br>Exitosas: %{x}<extra></extra>",
            ))
            fig_tec.add_trace(go.Bar(
                y=resumen_sorted["TECNICO"], x=resumen_sorted["Fallidas"],
                name="Fallidas", orientation="h",
                marker=dict(color=COLOR_DANGER, line_width=0),
                hovertemplate="<b>%{y}</b><br>Fallidas: %{x}<extra></extra>",
            ))
            fig_tec.update_layout(
                barmode="stack",
                title=dict(text="Visitas por técnico", font=dict(size=15, color="#333"), x=0.01),
                height=max(280, len(resumen_sorted) * 42 + 80),
                margin=dict(l=10, r=10, t=50, b=10),
                plot_bgcolor="white", paper_bgcolor="white",
                legend=dict(orientation="h", yanchor="bottom", y=1.01, xanchor="right", x=1, font=dict(size=11)),
                xaxis=dict(showgrid=True, gridcolor="#f5f5f5", title="Visitas", tickfont=dict(size=10)),
                yaxis=dict(showgrid=False, tickfont=dict(size=11)),
            )
            st.plotly_chart(fig_tec, use_container_width=True)

    # ── TAB POR PREDIO ────────────────────────────────────────
    with tab_pred:
        st.markdown('<div class="section-title">Historial por predio</div>', unsafe_allow_html=True)
        rea_txt = st.text_input("Buscar REA (escriba algunos caracteres)", placeholder="Ej: 000123")
        rea_busq = ""
        if rea_txt and "REA" in visitas.columns:
            opciones_rea = sorted(visitas["REA"].dropna().astype(str).unique().tolist())
            filtradas = [r for r in opciones_rea if rea_txt.strip().lower() in r.lower()]
            if filtradas:
                rea_busq = st.selectbox("Seleccione el predio", filtradas, key="sel_rea_hist")
            else:
                st.warning(f"No hay coincidencias para '{rea_txt}'")

        if rea_busq:
            vis_predio = visitas[visitas["REA"] == rea_busq] if "REA" in visitas.columns else pd.DataFrame()
            res_predio = resultados[resultados["REA"] == rea_busq] if "REA" in resultados.columns else pd.DataFrame()

            if vis_predio.empty:
                st.warning(f"No se encontraron visitas para REA '{rea_busq}'")
            else:
                maestro = load_maestro()
                if not maestro.empty:
                    predio_m = maestro[maestro["REA"] == rea_busq.strip()]
                    if not predio_m.empty:
                        pi = predio_m.iloc[0]
                        st.markdown(
                            f"""
                            **Predio:** {pi.get('REA','')} | **Dirección:** {pi.get('DIRECCION','')} |
                            **Barrio:** {pi.get('BARRIO','')} | **Localidad:** {pi.get('LOCALIDAD','')} |
                            **Estado REA:** {pi.get('ESTADO_REA','')}
                            """
                        )

                st.markdown(f"**Total visitas realizadas:** {len(vis_predio)}")

                show_vp = [c for c in ["NUM_VISITA", "FECHA_PROGRAMADA", "TECNICOS", "ESTADO", "NUM_VISITA_PREDIO", "OBSERVACIONES_PROG"] if c in vis_predio.columns]
                st.dataframe(vis_predio[show_vp], use_container_width=True, hide_index=True)

                if not res_predio.empty:
                    st.markdown('<div class="section-title">Resultados registrados</div>', unsafe_allow_html=True)
                    show_rp = [c for c in ["NUM_VISITA", "FECHA_VISITA", "RESULTADO", "OCUPACION", "TIPO_CONSTRUCCION", "OBSERVACIONES"] if c in res_predio.columns]
                    st.dataframe(res_predio[show_rp], use_container_width=True, hide_index=True)

                # Timeline
                if "FECHA_PROGRAMADA" in vis_predio.columns and "ESTADO" in vis_predio.columns:
                    st.markdown('<div class="section-title">Línea de tiempo</div>', unsafe_allow_html=True)
                    tl_df = vis_predio[["NUM_VISITA", "FECHA_PROGRAMADA", "ESTADO", "TECNICOS"]].copy()
                    tl_df["_fecha"] = pd.to_datetime(tl_df["FECHA_PROGRAMADA"], errors="coerce")
                    tl_df = tl_df.dropna(subset=["_fecha"])
                    if not tl_df.empty:
                        fig_tl = px.scatter(
                            tl_df,
                            x="_fecha",
                            y=[rea_busq] * len(tl_df),
                            color="ESTADO",
                            color_discrete_map={"Exitosa": COLOR_SUCCESS, "Fallida": COLOR_DANGER, "Pendiente": COLOR_WARNING},
                            hover_data=["NUM_VISITA", "TECNICOS"],
                            title=f"Línea de tiempo - REA {rea_busq}",
                            size_max=18,
                            size=[20] * len(tl_df),
                        )
                        fig_tl.update_layout(
                            yaxis_visible=False,
                            plot_bgcolor="white",
                            paper_bgcolor="white",
                        )
                        st.plotly_chart(fig_tl, use_container_width=True)

    # ── TAB REPORTE EXCEL ─────────────────────────────────────
    with tab_rep:
        st.markdown('<div class="section-title">Reporte de Gestión por Técnico</div>', unsafe_allow_html=True)
        st.caption("Seleccione el período y técnico para generar el reporte en Excel.")

        if visitas.empty:
            st.info("No hay datos de visitas aún.")
        else:
            rp1, rp2, rp3 = st.columns(3)
            with rp1:
                rp_fecha_ini = st.date_input("Fecha inicio", value=date.today().replace(day=1), key="rp_fi")
            with rp2:
                rp_fecha_fin = st.date_input("Fecha fin", value=date.today(), key="rp_ff")
            with rp3:
                tec_all = ["Todos"]
                if "TECNICOS" in visitas.columns:
                    all_t = set()
                    for t in visitas["TECNICOS"].dropna():
                        all_t.update(t.split("|"))
                    tec_all += sorted(all_t)
                rp_tecnico = st.selectbox("Técnico", tec_all, key="rp_tec")

            if st.button("Generar Reporte", use_container_width=False, key="btn_reporte"):
                v_rep = visitas.copy()
                v_rep["_fd"] = pd.to_datetime(v_rep.get("FECHA_PROGRAMADA"), errors="coerce")
                v_rep = v_rep[
                    (v_rep["_fd"] >= pd.Timestamp(rp_fecha_ini))
                    & (v_rep["_fd"] <= pd.Timestamp(rp_fecha_fin))
                ]

                # Explode by técnico
                v_rep_exp = v_rep.assign(
                    TECNICO=v_rep["TECNICOS"].str.split("|")
                ).explode("TECNICO")
                v_rep_exp["TECNICO"] = v_rep_exp["TECNICO"].str.strip()
                if rp_tecnico != "Todos":
                    v_rep_exp = v_rep_exp[v_rep_exp["TECNICO"] == rp_tecnico]

                if v_rep_exp.empty:
                    st.warning("No hay visitas en el período/técnico seleccionado.")
                else:
                    # Compute rendimiento semanal (promedio último mes dentro del rango)
                    cutoff_rend = pd.Timestamp(rp_fecha_fin) - pd.Timedelta(weeks=4)
                    v_rend_base = v_rep_exp[v_rep_exp["_fd"] >= cutoff_rend].copy()

                    resumen_rep = []
                    for tec in sorted(v_rep_exp["TECNICO"].dropna().unique()):
                        df_t = v_rep_exp[v_rep_exp["TECNICO"] == tec]
                        programadas = len(df_t)
                        exitosas_t  = int((df_t.get("ESTADO", pd.Series()) == "Exitosa").sum())
                        fallidas_t  = int((df_t.get("ESTADO", pd.Series()) == "Fallida").sum())
                        realizadas_t = exitosas_t + fallidas_t
                        pct_t = round(exitosas_t / realizadas_t * 100, 1) if realizadas_t > 0 else 0
                        # rendimiento semanal (último mes)
                        df_t_rend = v_rend_base[v_rend_base["TECNICO"] == tec]
                        sems_rend = df_t_rend["_fd"].dt.to_period("W").nunique() if not df_t_rend.empty else 0
                        rend_t = round(len(df_t_rend) / max(sems_rend, 1), 1) if not df_t_rend.empty else 0.0
                        resumen_rep.append({
                            "Técnico": tec,
                            "Visitas Programadas": programadas,
                            "Visitas Exitosas": exitosas_t,
                            "Visitas Fallidas": fallidas_t,
                            "% Éxito": f"{pct_t}%",
                            "Rendimiento Semanal (prom. 1 mes)": rend_t,
                        })

                    df_rep = pd.DataFrame(resumen_rep)
                    st.dataframe(df_rep, use_container_width=True, hide_index=True)

                    buf_rep = io.BytesIO()
                    with pd.ExcelWriter(buf_rep, engine="openpyxl") as writer:
                        df_rep.to_excel(writer, index=False, sheet_name="Reporte")
                    st.download_button(
                        label="⬇️ Descargar Reporte Excel",
                        data=buf_rep.getvalue(),
                        file_name=f"ReporteGestion_{rp_fecha_ini}_{rp_fecha_fin}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=False,
                        key="dl_reporte",
                    )

